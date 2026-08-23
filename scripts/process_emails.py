#!/usr/bin/env python3
"""
Email ingestion script for FamilyBot.
Polls the familybot Gmail inbox, processes new emails, parses attachments,
and stores results in the family SQLite database.
"""

import subprocess
import json
import sqlite3
import os
import sys
import tempfile
import re
import html
from datetime import datetime
sys.path.insert(0, os.path.dirname(__file__))
from email_routing import resolve_sender, assert_no_member_mix, get_quarantined, _resolve_member_from_body
from parse_ukeplan import parse_ukeplan_pdf
from reliability import (
    EmailProcessingLedger,
    atomic_write_json,
    connect_db,
    database_path,
    workspace_path,
)

DB_PATH = str(database_path())
ATTACHMENT_DIR = str(workspace_path() / "db" / "attachments")
os.makedirs(ATTACHMENT_DIR, exist_ok=True)

def run(cmd, timeout=60):
    """Run one command without a shell so message IDs cannot become code."""
    env = os.environ.copy()
    return subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        env=env,
        timeout=timeout,
    )

def get_db():
    return connect_db(DB_PATH)

def already_processed(message_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM email_log WHERE message_id = ?", (str(message_id),))
    row = c.fetchone()
    conn.close()
    return row is not None

def log_email(message_id, subject, sender, received_at, category, has_pdf, summary, member_id=None):
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        INSERT OR IGNORE INTO email_log
            (message_id, subject, sender, received_at, processed_at, member_id, category, has_pdf, summary)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (str(message_id), subject, sender, received_at,
          datetime.now().isoformat(), member_id, category, has_pdf, summary))
    conn.commit()
    conn.close()

def save_ukeplan(pdf_path, member_id, email_id, year=None):
    """Parse a ukeplan PDF and store it in week_plans + week_plan_days."""
    if year is None:
        year = datetime.now().year
    try:
        data = parse_ukeplan_pdf(pdf_path, member_id=member_id, year=year)
    except Exception as e:
        print(f"  [UKEPLAN] Parse error: {e}")
        return False

    week = data.get("week")
    if not week:
        print(f"  [UKEPLAN] Could not determine week number from {pdf_path}")
        return False

    summary_parts = []
    if data.get("theme"):
        summary_parts.append(f"Tema: {data['theme']}")
    if data.get("homework"):
        hw = ", ".join(f"{k}: {v}" for k, v in data["homework"].items())
        summary_parts.append(f"Lekser: {hw}")
    if data.get("info"):
        summary_parts.append(data["info"][:300])
    summary = " | ".join(summary_parts)
    raw_text = clean_text(data.get("raw_text") or data.get("info") or "")
    if not summary:
        summary = raw_text[:900]

    conn = get_db()
    c = conn.cursor()
    teacher = data.get("teacher") or ""
    if not teacher:
        teacher_row = c.execute("SELECT teacher FROM family_members WHERE id=?", (member_id,)).fetchone()
        teacher = (teacher_row[0] if teacher_row else "") or ""

    row = c.execute("""
        SELECT id, source_email_id, raw_text FROM week_plans
        WHERE member_id=? AND week_number=? AND year=?
        ORDER BY created_at DESC, id DESC LIMIT 1
    """, (member_id, week, year)).fetchone()
    if row:
        plan_id, existing_source, existing_raw = row
        try:
            if existing_source and int(existing_source) > int(email_id):
                conn.close()
                return False
        except (TypeError, ValueError):
            pass
        if existing_source == str(email_id) and (existing_raw or "") == raw_text:
            conn.close()
            return True
        c.execute("""
            UPDATE week_plans
            SET raw_text=?, summary=?, source_email_id=?, teacher=?, created_at=datetime('now')
            WHERE id=?
        """, (raw_text, summary, str(email_id), teacher, plan_id))
        c.execute("DELETE FROM week_plan_days WHERE week_plan_id=?", (plan_id,))
    else:
        c.execute("""
            INSERT INTO week_plans (member_id, week_number, year, raw_text, summary, source_email_id, teacher)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (member_id, week, year, raw_text, summary, str(email_id), teacher))
        plan_id = c.lastrowid

    for iso, day_data in data.get("days", {}).items():
        weekday = day_data.get("weekday", "")
        events = "; ".join(day_data.get("events", []))
        bring = "; ".join(day_data.get("bring", []))
        hw_day = "; ".join(f"{k}: {v}" for k, v in data.get("homework", {}).items())
        c.execute("""
            INSERT INTO week_plan_days (week_plan_id, day, date, note, homework, bring)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (plan_id, weekday, iso, events, hw_day, bring))

    conn.commit()
    conn.close()
    print(f"  [UKEPLAN] Saved week {week}/{year} for member {member_id} (plan_id={plan_id})")
    return bool(plan_id)


def clean_text(text: str) -> str:
    text = html.unescape(text or "")
    text = re.sub(r"<#part[^>]*>", " ", text)
    text = re.sub(r"<#/part>", " ", text)
    text = re.sub(r"Dette er en autogenerert e-post fra Portalen.*", "", text, flags=re.I | re.S)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def strip_email_metadata(text: str) -> str:
    """Remove transport headers, attachment markers and addresses from text."""
    text = clean_text(text)
    attachment = re.search(r"\[[^\]]*attachment:\s*[^\]]+\]", text, flags=re.IGNORECASE)
    if attachment:
        text = text[attachment.end():].lstrip()
    else:
        subject = re.search(r"\bSubject:\s*", text, flags=re.IGNORECASE)
        prefix = text[:subject.start()] if subject else ""
        if subject and re.search(r"\b(?:From|To|Cc|Bcc|Date|Reply-To):", prefix, flags=re.IGNORECASE):
            text = text[subject.end():].lstrip()
    text = re.sub(r"(?im)^\s*(?:from|to|cc|bcc|date|reply-to|subject):[^\n]*\n?", "", text)
    text = re.sub(r"\b(?:from|to|cc|bcc|date|reply-to|subject):\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b", "", text)
    text = re.sub(r"\s+[A-Za-z][\w .-]*\s*<\s*>", " ", text)
    text = re.sub(r"\[[^\]]*attachment:[^\]]*\]", "", text, flags=re.IGNORECASE)
    return re.sub(r"\s{2,}", " ", text).strip()


def declared_attachment_names(message_text: str) -> set[str]:
    names = set()
    for match in re.finditer(r"<#part\b[^>]*\bfilename=\"([^\"]+)\"", message_text or ""):
        names.add(os.path.basename(html.unescape(match.group(1))))
    return names


def save_ukeplan_attachments(paths, member_id, email_id, subject="", year=None):
    """Store a ukeplan from document attachments, never from the email body.

    A PDF is authoritative when present. Falling back to ``full_text`` after a
    PDF parse failure used to persist mail headers, addresses and unrelated
    prose as if it were the plan. That is both noisy and a data-boundary bug;
    the caller should leave the message retryable when the PDF cannot be
    parsed.
    """
    paths = [path for path in paths if path]
    pdf_paths = [path for path in paths if path.lower().endswith(".pdf")]
    if pdf_paths:
        saved = False
        for path in pdf_paths:
            saved = save_ukeplan(path, member_id, email_id, year=year) or saved
        if not saved:
            raise RuntimeError("ukeplan PDF was present but could not be parsed")
        return True

    docx_paths = [path for path in paths if path.lower().endswith(".docx")]
    for path in docx_paths:
        if save_ukeplan_text(extract_docx(path), member_id, email_id, subject, year=year):
            return True
    return False


def save_ukeplan_text(raw_text, member_id, email_id, subject="", year=None):
    """Store a ukeplan when the school message body has the plan but no PDF parser output."""
    if year is None:
        year = datetime.now().year

    text = clean_text(raw_text)
    m = re.search(r"uke\s*(\d+)", f"{subject} {text}", re.I)
    if not m:
        return False
    week = int(m.group(1))

    teacher = ""
    m = re.search(r"Av:\s*([^,]+),\s*Dato:", text, re.I)
    if m:
        teacher = m.group(1).strip()

    summary = text
    # Keep the useful part compact; the LLM gets raw_text too when needed.
    for marker in ["Hei alle sammen", "Hei", "Mandag"]:
        idx = summary.lower().find(marker.lower())
        if idx >= 0:
            summary = summary[idx:]
            break
    summary = summary[:900].strip()

    conn = get_db()
    c = conn.cursor()
    row = c.execute("""
        SELECT id, raw_text, source_email_id FROM week_plans
        WHERE member_id=? AND week_number=? AND year=?
        ORDER BY created_at DESC LIMIT 1
    """, (member_id, week, year)).fetchone()

    if row:
        plan_id, existing_raw, existing_source = row
        try:
            if existing_source and int(existing_source) > int(email_id):
                conn.close()
                return False
        except ValueError:
            pass
        if existing_source == str(email_id) and (existing_raw or "") == text:
            conn.close()
            return False
        c.execute("""
            UPDATE week_plans
            SET raw_text=?, summary=?, source_email_id=?, teacher=?
            WHERE id=?
        """, (text, summary, str(email_id), teacher, plan_id))
    else:
        c.execute("""
            INSERT INTO week_plans (member_id, week_number, year, raw_text, summary, source_email_id, teacher)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (member_id, week, year, text, summary, str(email_id), teacher))
        plan_id = c.lastrowid

    conn.commit()
    conn.close()
    print(f"  [UKEPLAN] Saved text week {week}/{year} for member {member_id} (plan_id={plan_id})")
    return True


def backfill_weekplans_from_email_log():
    """Recover recent ukeplans stranded before a class mapping update."""
    conn = get_db()
    rows = conn.execute("""
        SELECT message_id, subject, member_id, summary
        FROM email_log
        WHERE category='school'
          AND member_id IS NULL
          AND lower(subject) LIKE '%ukeplan%'
          AND date(processed_at) >= date('now', '-45 days')
        ORDER BY id
    """).fetchall()
    conn.close()

    count = 0
    for message_id, subject, member_id, summary in rows:
        resolved_member = member_id or _resolve_member_from_body(f"{subject} {summary or ''}")
        if not resolved_member:
            continue

        conn = get_db()
        plan_exists = conn.execute(
            "SELECT 1 FROM week_plans WHERE source_email_id=? LIMIT 1", (str(message_id),)
        ).fetchone()
        if plan_exists:
            conn.close()
            continue
        conn.execute("UPDATE email_log SET member_id=? WHERE message_id=?", (resolved_member, str(message_id)))
        conn.execute("""
            UPDATE routing_audit
            SET resolved_member_id=?, notes=notes || ' Backfilled after class mapping update.'
            WHERE message_id=? AND resolved_member_id IS NULL
        """, (resolved_member, str(message_id)))
        conn.commit()
        conn.close()

        attachment_paths = []
        for name in declared_attachment_names(summary or ""):
            path = os.path.join(ATTACHMENT_DIR, os.path.basename(name))
            if not os.path.isfile(path):
                continue
            attachment_paths.append(path)
        try:
            saved = save_ukeplan_attachments(
                attachment_paths, resolved_member, message_id, subject=subject
            )
        except RuntimeError as exc:
            print(f"  [UKEPLAN] {exc}")
            saved = False
        if saved:
            count += 1
    return count


def extract_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[docx parse error: {e}]"

def extract_pdf(path):
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
                # Try table extraction too
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if cells:
                            text.append(" | ".join(cells))
        return "\n".join(text)
    except Exception as e:
        return f"[pdf parse error: {e}]"

def process_new_emails():
    ledger = EmailProcessingLedger(DB_PATH)
    # List unread emails as JSON
    listed = run(["himalaya", "envelope", "list", "--output", "json", "--page-size", "20"])
    out, err = listed.stdout.strip(), listed.stderr.strip()
    if listed.returncode != 0:
        raise RuntimeError(f"himalaya envelope list failed: {(err or out)[:300]}")
    if not out:
        print(f"No output from himalaya or inbox empty. stderr={err[:300]}")
        return []

    try:
        envelopes = json.loads(out)
    except json.JSONDecodeError:
        print("Failed to parse himalaya output:", out[:200])
        return []

    new_items = []
    failures = []

    for env in envelopes:
        msg_id = str(env.get("id", ""))
        subject = env.get("subject", "(no subject)")
        sender = env.get("from", {})
        if isinstance(sender, dict):
            sender_str = sender.get("addr", "") or sender.get("name", "")
        else:
            sender_str = str(sender)
        date_str = env.get("date", "")

        ledger_status = ledger.status(msg_id)
        if ledger_status in ledger.TERMINAL_STATES:
            continue
        # Existing installations have email_log rows from before the ledger.
        # Only use that as a terminal signal when no ledger record exists;
        # a failed ledger row must remain retryable even if logging succeeded.
        if ledger_status is None and already_processed(msg_id):
            continue

        print(f"Processing: [{msg_id}] {subject}")
        ledger.begin(msg_id, subject)

        try:
            # Read full message
            read = run(["himalaya", "message", "read", msg_id, "--output", "plain"])
            body_out, read_err = read.stdout.strip(), read.stderr.strip()
            if read.returncode != 0:
                raise RuntimeError(f"himalaya message read failed: {(read_err or body_out)[:300]}")
            ledger.update(msg_id, "message-read")

            before_files = set(os.listdir(ATTACHMENT_DIR))

            # Download attachments. A failed download is retryable when the
            # message declares attachments; silently processing partial input
            # would make the email look complete forever.
            attachment_names = declared_attachment_names(body_out)
            downloaded = run([
                "himalaya", "attachment", "download",
                "--downloads-dir", ATTACHMENT_DIR,
                msg_id,
            ])
            if downloaded.returncode != 0 and attachment_names:
                detail = (downloaded.stderr or downloaded.stdout).strip()
                raise RuntimeError(f"attachment download failed: {detail[:300]}")
            if downloaded.returncode != 0:
                print(f"  [WARN] attachment check failed: {(downloaded.stderr or downloaded.stdout)[:200]}")
            ledger.update(msg_id, "attachments-read")

            # Find attachments declared in the message or created by the download.
            attachment_text = ""
            has_pdf = False
            fresh_ukeplan_files = []
            for fname in os.listdir(ATTACHMENT_DIR):
                fpath = os.path.join(ATTACHMENT_DIR, fname)
                if fname not in before_files or fname in attachment_names:
                    if fname.lower().endswith(".docx"):
                        attachment_text += f"\n[Attachment: {fname}]\n" + extract_docx(fpath)
                        fresh_ukeplan_files.append(fpath)
                    elif fname.lower().endswith(".pdf"):
                        has_pdf = True
                        attachment_text += f"\n[Attachment: {fname}]\n" + extract_pdf(fpath)
                        fresh_ukeplan_files.append(fpath)

            full_text = body_out + attachment_text

            # Route through guard rails — always, no exceptions
            routing = resolve_sender(msg_id, subject, sender_str,
                                      env.get("from", {}).get("name", sender_str) if isinstance(env.get("from"), dict) else sender_str,
                                      full_text)
            category = routing.category
            member_id = routing.member_id
            ledger.update(msg_id, "routed")

            if routing.quarantine:
                print(f"  [QUARANTINE] {subject} from {sender_str} — {routing.notes}")

            log_email(msg_id, subject, sender_str, date_str, category, int(has_pdf), full_text[:2000], member_id)
            ledger.update(msg_id, "email-logged")

            # A ukeplan is only terminal after it has produced a week_plan row.
            if category == "school" and member_id:
                ukeplan_keywords = ["ukeplan", "week plan", "ukeplanen"]
                if any(kw in subject.lower() for kw in ukeplan_keywords):
                    attachment_paths = []
                    likely_plans = [p for p in fresh_ukeplan_files if "ukeplan" in os.path.basename(p).lower()]
                    candidates = likely_plans or fresh_ukeplan_files
                    attachment_paths.extend(candidates)
                    try:
                        saved = save_ukeplan_attachments(
                            attachment_paths, member_id, msg_id, subject=subject
                        )
                    except RuntimeError as exc:
                        print(f"  [UKEPLAN] {exc}")
                        saved = False
                    if not saved:
                        raise RuntimeError("ukeplan detected but no week plan was stored")
                    ledger.update(msg_id, "ukeplan-stored")

            new_items.append({
                "id": msg_id,
                "subject": subject,
                "sender": sender_str,
                "category": category,
                "member_id": member_id,
                "has_attachment": bool(fresh_ukeplan_files),
                "text": full_text
            })
            ledger.complete(msg_id, quarantined=routing.quarantine)

            # Mark as read only after durable completion.
            flagged = run(["himalaya", "flag", "add", msg_id, "seen"])
            if flagged.returncode != 0:
                print(f"  [WARN] could not mark {msg_id} seen: {(flagged.stderr or flagged.stdout)[:200]}")
        except Exception as exc:
            error = f"{type(exc).__name__}: {exc}"
            try:
                ledger.fail(msg_id, "failed", error)
            except Exception:
                pass
            failures.append((msg_id, error))
            print(f"  [FAILED] {msg_id}: {error}")

    if failures:
        detail = "; ".join(f"{message_id}: {error}" for message_id, error in failures)
        raise RuntimeError(f"{len(failures)} email(s) failed and remain retryable: {detail[:1000]}")
    return new_items

if __name__ == "__main__":
    print(f"[{datetime.now().isoformat()}] Checking email...")
    items = process_new_emails()
    if items:
        print(f"Processed {len(items)} new email(s):")
        for item in items:
            print(f"  - [{item['category']}] {item['subject']} (member: {item['member_id']}, attachment: {item['has_attachment']})")
        # Write summary for the agent to pick up
        summary_path = workspace_path() / "db" / "new_emails.json"
        atomic_write_json(summary_path, items)
        print(f"Summary written to {summary_path}")
    else:
        print("No new emails.")
    backfilled = backfill_weekplans_from_email_log()
    if backfilled:
        print(f"Backfilled/updated {backfilled} ukeplan record(s) from email_log.")
